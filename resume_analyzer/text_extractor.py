# -*- coding: utf-8 -*-
"""Functions for extracting text content from various file formats."""

import os
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    text = ""
    if not os.path.exists(pdf_path):
        print(f"Error: PDF file not found at {pdf_path}")
        return text
    try:
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        # Catch specific PyPDF2 exceptions if needed, e.g., PasswordError
        print(f"Error reading PDF {pdf_path}: {e}")
    return text

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    text = ""
    if not os.path.exists(docx_path):
        print(f"Error: DOCX file not found at {docx_path}")
        return text
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # TODO: Add table extraction from DOCX if required
        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             text += cell.text + "\t" # Or some other separator
        #         text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
    return text

def extract_text_from_txt(txt_path):
    """Extracts text from a TXT file."""
    text = ""
    if not os.path.exists(txt_path):
        print(f"Error: TXT file not found at {txt_path}")
        return text
    try:
        # Try common encodings
        encodings_to_try = ["utf-8", "latin-1", "cp1252"]
        for enc in encodings_to_try:
            try:
                with open(txt_path, 'r', encoding=enc) as f:
                    text = f.read()
                break # Stop if successful
            except UnicodeDecodeError:
                continue # Try next encoding
        else:
             print(f"Warning: Could not decode TXT file {txt_path} with common encodings.")
             # Optionally, try reading with errors ignored
             # with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
             #     text = f.read()

    except Exception as e:
        print(f"Error reading TXT {txt_path}: {e}")
    return text

def extract_text(file_path):
    """Extracts text from PDF, DOCX, or TXT files based on extension."""
    if not file_path or not isinstance(file_path, str):
        print("Error: Invalid file path provided.")
        return ""

    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return ""

    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    if file_extension == ".pdf":
        print(f"Extracting text from PDF: {file_path}")
        return extract_text_from_pdf(file_path)
    elif file_extension == ".docx":
        print(f"Extracting text from DOCX: {file_path}")
        return extract_text_from_docx(file_path)
    elif file_extension == ".txt":
        print(f"Extracting text from TXT: {file_path}")
        return extract_text_from_txt(file_path)
    else:
        print(f"Unsupported file format: {file_extension}. Only PDF, DOCX, TXT are supported.")
        return ""

